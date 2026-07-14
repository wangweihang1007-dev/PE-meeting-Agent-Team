import argparse
import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


IMAGE_PATTERN = re.compile(r"^!\[(?P<alt>.*?)\]\((?P<path>.*?)\)$")
QUESTION_LINE_PATTERN = re.compile(r"^(?P<label>(?:Q|问题)\s*[:：])(?P<body>.*)$", re.IGNORECASE)


def has_style(document: Document, style_name: str) -> bool:
    try:
        document.styles[style_name]
        return True
    except KeyError:
        return False


def safe_style(document: Document, preferred: str, fallback: str = "Normal") -> str:
    return preferred if has_style(document, preferred) else fallback


def clear_body(document: Document) -> None:
    body = document._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def add_run_with_inline_bold(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def add_paragraph(document: Document, text: str, style: str = "Normal"):
    paragraph = document.add_paragraph(style=safe_style(document, style))
    add_run_with_inline_bold(paragraph, text)
    return paragraph


def add_plain_heading(document: Document, text: str, level: int):
    paragraph = document.add_paragraph(style=safe_style(document, "Normal"))
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.space_before = Pt(8 if level == 1 else 6)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt({1: 16, 2: 14, 3: 12}.get(level, 12))
    return paragraph


def apply_body_indent(paragraph) -> None:
    paragraph.paragraph_format.first_line_indent = Pt(21)


def flush_body_list(document: Document, items: list[str]) -> None:
    if not items:
        return
    paragraph = add_paragraph(document, "；".join(item.rstrip("。；;") for item in items) + "。")
    apply_body_indent(paragraph)
    items.clear()


def normalize_markdown_line(line: str) -> str:
    return line.strip().replace("\\_", "_")


def add_markdown_image(document: Document, markdown_path: Path, line: str) -> bool:
    match = IMAGE_PATTERN.match(line)
    if not match:
        return False
    image_path = Path(match.group("path").strip().strip('"'))
    if not image_path.is_absolute():
        image_path = markdown_path.parent / image_path
    if not image_path.exists():
        add_paragraph(document, f"[图片缺失：{match.group('alt') or image_path.name}] {image_path}", style="Normal")
        return True

    paragraph = document.add_paragraph(style=safe_style(document, "Normal"))
    paragraph.paragraph_format.first_line_indent = None
    try:
        paragraph.add_run().add_picture(str(image_path), width=Inches(6.0))
    except Exception as exc:
        paragraph.add_run(f"[图片插入失败：{image_path.name}；{exc}]")
    alt = match.group("alt").strip()
    if alt:
        caption = add_paragraph(document, alt, style="Normal")
        caption.paragraph_format.first_line_indent = None
    return True


def heading_style_for_line(line: str) -> tuple[str, int] | None:
    if line.startswith("#"):
        level = len(line) - len(line.lstrip("#"))
        text = line[level:].strip()
        text = re.sub(r"^\d+(?:\.\d+)*\s+", "", text)
        return text, min(level, 3)
    if re.match(r"^\d+\.\d+\s+.+", line):
        return re.sub(r"^\d+\.\d+\s+", "", line), 2
    if re.match(r"^\d+\s+.+", line):
        return re.sub(r"^\d+\s+", "", line), 1
    if line.startswith("附录"):
        return line, 1
    return None


def markdown_to_docx(markdown_path: Path, template_path: Path, output_path: Path) -> None:
    document = Document(str(template_path))
    clear_body(document)
    text = markdown_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    title = document.add_paragraph(style="Title")
    title.add_run("投资访谈会议纪要")
    in_appendix = False
    body_list_items: list[str] = []

    for raw_line in lines:
        line = normalize_markdown_line(raw_line)
        if not line:
            continue
        if line.startswith("【分区") or line.startswith("【时间范围"):
            continue
        if set(line) <= {"-"}:
            continue
        if IMAGE_PATTERN.match(line):
            flush_body_list(document, body_list_items)
            add_markdown_image(document, markdown_path, line)
            continue
        heading = heading_style_for_line(line)
        if heading:
            flush_body_list(document, body_list_items)
            heading_text, level = heading
            if heading_text == "会议纪要":
                continue
            if heading_text.startswith("附录"):
                in_appendix = True
            add_plain_heading(document, heading_text, level=level)
            continue
        if line.startswith("- "):
            if in_appendix:
                flush_body_list(document, body_list_items)
                style = "List Bullet" if has_style(document, "List Bullet") else "List Paragraph"
                text = line[2:] if has_style(document, "List Bullet") else "• " + line[2:]
                add_paragraph(document, text, style=style)
            else:
                body_list_items.append(line[2:])
            continue
        if re.match(r"^\d+\.\s+", line):
            flush_body_list(document, body_list_items)
            if in_appendix and has_style(document, "List Number"):
                add_paragraph(document, line, style="List Number")
            else:
                paragraph = add_paragraph(document, line, style="Normal")
                apply_body_indent(paragraph)
            continue
        question_match = QUESTION_LINE_PATTERN.match(line)
        if question_match:
            flush_body_list(document, body_list_items)
            paragraph = document.add_paragraph(style="Normal")
            run = paragraph.add_run(f"{question_match.group('label')}{question_match.group('body')}")
            run.bold = True
            continue
        if line.startswith("A："):
            flush_body_list(document, body_list_items)
            paragraph = document.add_paragraph(style="Normal")
            run = paragraph.add_run("A：")
            run.bold = True
            paragraph.add_run(line[2:])
            continue
        flush_body_list(document, body_list_items)
        paragraph = add_paragraph(document, line)
        if not in_appendix:
            apply_body_indent(paragraph)

    flush_body_list(document, body_list_items)

    for section in document.sections:
        section.top_margin = section.top_margin
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))


def main() -> None:
    parser = argparse.ArgumentParser(description="使用指定 Word 模板把 Markdown 纪要转换为 DOCX")
    parser.add_argument("--markdown", required=True)
    parser.add_argument("--template", required=True)
    parser.add_argument("--output", required=True)
    args = parser.parse_args()
    markdown_to_docx(Path(args.markdown), Path(args.template), Path(args.output))
    print(f"完成。输出文件：{Path(args.output).resolve()}")


if __name__ == "__main__":
    main()
